# services/report_generator.py - Word Report Generator
import os
from datetime import datetime
from typing import List, Dict, Any
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE


class ReportGenerator:
    """Generate professional Word reports for legal risk analysis"""

    def __init__(self):
        self.doc = None

    def _add_styles(self):
        """Add custom styles to the document"""
        styles = self.doc.styles

        # Modify existing styles for better appearance
        heading1 = styles['Heading 1']
        heading1.font.size = Pt(16)
        heading1.font.bold = True
        heading1.font.color.rgb = RGBColor(0, 51, 102)

        heading2 = styles['Heading 2']
        heading2.font.size = Pt(14)
        heading2.font.bold = True
        heading2.font.color.rgb = RGBColor(0, 51, 102)

        heading3 = styles['Heading 3']
        heading3.font.size = Pt(12)
        heading3.font.bold = True

    def _add_cover_page(self, analysis_name: str, analysis_id: str):
        """Add cover page to the report"""
        # Add title
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("\n\n\n\nLEGAL RISK ANALYSIS REPORT")
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)

        # Add analysis name
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f"\n{analysis_name}")
        run.font.size = Pt(16)

        # Add metadata
        meta = self.doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = meta.add_run(f"\n\nAnalysis ID: {analysis_id}\nDate: {datetime.now().strftime('%B %d, %Y')}")
        run.font.size = Pt(12)

        # Add confidentiality notice
        notice = self.doc.add_paragraph()
        notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = notice.add_run("\n\n\n\nCONFIDENTIAL - ATTORNEY WORK PRODUCT")
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(128, 0, 0)

        self.doc.add_page_break()

    def _add_executive_summary(self, analysis_summary: str, risks: List[Dict]):
        """Add executive summary section"""
        self.doc.add_heading('Executive Summary', level=1)

        # Overall assessment
        self.doc.add_paragraph(analysis_summary)

        # Risk counts by severity
        severity_counts = {"Critical": 0, "High": 0, "Medium": 0, "Low": 0}
        for risk in risks:
            severity = risk.get("severity", "Medium")
            if severity in severity_counts:
                severity_counts[severity] += 1

        # Summary table
        table = self.doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'

        # Header
        cells = table.rows[0].cells
        cells[0].text = "Severity Level"
        cells[1].text = "Count"
        for cell in cells:
            cell.paragraphs[0].runs[0].font.bold = True

        # Data rows
        for i, (severity, count) in enumerate(severity_counts.items(), 1):
            cells = table.rows[i].cells
            cells[0].text = severity
            cells[1].text = str(count)

        self.doc.add_paragraph()

        # Top risks
        critical_high = [r for r in risks if r.get("severity") in ["Critical", "High"]]
        if critical_high:
            self.doc.add_heading('Top Priority Risks', level=2)
            for risk in critical_high[:5]:
                para = self.doc.add_paragraph(style='List Bullet')
                run = para.add_run(f"{risk.get('title', 'Untitled Risk')} ")
                run.font.bold = True
                para.add_run(f"({risk.get('severity', 'Unknown')} - {risk.get('category', 'Unknown')})")

        self.doc.add_page_break()

    def _add_methodology(self, documents: List[Dict]):
        """Add methodology section"""
        self.doc.add_heading('Methodology', level=1)

        self.doc.add_paragraph(
            "This legal risk analysis was conducted using a systematic review of all documents "
            "provided in the data room. Each document was analyzed for potential legal risks "
            "across five key categories: Contractual, Regulatory, Litigation, Intellectual Property, "
            "and Operational risks."
        )

        self.doc.add_heading('Documents Reviewed', level=2)

        # Documents table
        table = self.doc.add_table(rows=len(documents) + 1, cols=3)
        table.style = 'Table Grid'

        # Header
        header_cells = table.rows[0].cells
        header_cells[0].text = "Document ID"
        header_cells[1].text = "Filename"
        header_cells[2].text = "Pages"
        for cell in header_cells:
            cell.paragraphs[0].runs[0].font.bold = True

        # Data
        for i, doc in enumerate(documents, 1):
            cells = table.rows[i].cells
            cells[0].text = doc.get('doc_id', '')
            cells[1].text = doc.get('original_filename', '')
            cells[2].text = str(doc.get('page_count', 0))

        self.doc.add_paragraph()
        self.doc.add_paragraph(f"Analysis Date: {datetime.now().strftime('%B %d, %Y')}")

        self.doc.add_page_break()

    def _add_detailed_risks(self, risks: List[Dict]):
        """Add detailed risk analysis section"""
        self.doc.add_heading('Detailed Risk Analysis', level=1)

        # Group risks by category
        categories = ["Contractual", "Regulatory", "Litigation", "IP", "Operational"]
        risks_by_category = {cat: [] for cat in categories}

        for risk in risks:
            category = risk.get("category", "Operational")
            if category in risks_by_category:
                risks_by_category[category].append(risk)

        # Add each category
        for category in categories:
            category_risks = risks_by_category[category]
            if not category_risks:
                continue

            self.doc.add_heading(f'{category} Risks', level=2)

            for risk in category_risks:
                # Risk title
                self.doc.add_heading(risk.get('title', 'Untitled Risk'), level=3)

                # Risk metadata
                meta_para = self.doc.add_paragraph()
                meta_para.add_run(f"Risk ID: ").font.bold = True
                meta_para.add_run(f"{risk.get('risk_id', 'N/A')}  |  ")
                meta_para.add_run(f"Severity: ").font.bold = True
                meta_para.add_run(f"{risk.get('severity', 'N/A')}  |  ")
                meta_para.add_run(f"Likelihood: ").font.bold = True
                meta_para.add_run(f"{risk.get('likelihood', 'N/A')}")

                # Description
                self.doc.add_heading('Description', level=4)
                self.doc.add_paragraph(risk.get('description', 'No description provided.'))

                # Evidence
                evidence = risk.get('evidence', [])
                if evidence:
                    self.doc.add_heading('Evidence', level=4)
                    for ev in evidence:
                        para = self.doc.add_paragraph()
                        para.add_run(f"Document: ").font.bold = True
                        para.add_run(f"{ev.get('doc_id', 'N/A')}, Page {ev.get('page_num', 'N/A')}")
                        quote_para = self.doc.add_paragraph()
                        run = quote_para.add_run(f'"{ev.get("citation", "")}"')
                        run.font.italic = True

                # Legal basis
                legal_basis = risk.get('legal_basis', '')
                if legal_basis:
                    self.doc.add_heading('Legal Basis', level=4)
                    self.doc.add_paragraph(legal_basis)

                # Mitigation
                mitigation = risk.get('recommended_mitigation', '')
                if mitigation:
                    self.doc.add_heading('Recommended Mitigation', level=4)
                    self.doc.add_paragraph(mitigation)

                self.doc.add_paragraph()  # Spacing between risks

        self.doc.add_page_break()

    def _add_recommendations(self, risks: List[Dict]):
        """Add recommendations summary section"""
        self.doc.add_heading('Summary of Recommendations', level=1)

        # Group by severity
        severity_order = ["Critical", "High", "Medium", "Low"]
        severity_labels = {
            "Critical": "Immediate Actions Required",
            "High": "Near-Term Actions",
            "Medium": "Medium-Term Improvements",
            "Low": "Monitoring Items"
        }

        for severity in severity_order:
            severity_risks = [r for r in risks if r.get("severity") == severity]
            if not severity_risks:
                continue

            self.doc.add_heading(severity_labels[severity], level=2)

            for risk in severity_risks:
                para = self.doc.add_paragraph(style='List Number')
                run = para.add_run(f"{risk.get('title', 'Untitled')}: ")
                run.font.bold = True
                para.add_run(risk.get('recommended_mitigation', 'No mitigation specified.'))

    def generate_report(
        self,
        analysis_id: str,
        analysis_name: str,
        analysis_summary: str,
        documents: List[Dict],
        risks: List[Dict],
        output_path: str
    ) -> str:
        """
        Generate a complete Word report

        Args:
            analysis_id: Unique analysis identifier
            analysis_name: Name of the analysis
            analysis_summary: Executive summary text
            documents: List of analyzed documents
            risks: List of identified risks
            output_path: Path to save the report

        Returns:
            Path to the generated report
        """
        # Create document
        self.doc = Document()
        self._add_styles()

        # Add sections
        self._add_cover_page(analysis_name, analysis_id)
        self._add_executive_summary(analysis_summary, risks)
        self._add_methodology(documents)
        self._add_detailed_risks(risks)
        self._add_recommendations(risks)

        # Ensure output directory exists
        os.makedirs(os.path.dirname(output_path), exist_ok=True)

        # Save document
        self.doc.save(output_path)

        return output_path


def create_report_generator() -> ReportGenerator:
    """Factory function to create the report generator"""
    return ReportGenerator()
